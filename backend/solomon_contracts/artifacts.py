"""
DOCX artifact generation for Solomon Contracts.
§9.1 Risk note (auto, bullet list by document)
§9.3 Protocol (4-column counterparty-facing table)
§9.2 Legal opinion (Sonnet markdown → DOCX)
"""
import io
import json
import logging
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def _safe_jsonb(value, default=None):
    """Parse a JSONB column that may already be a Python object (psycopg2 auto-parses JSONB)."""
    if default is None:
        default = []
    if value is None:
        return default
    if isinstance(value, (list, dict)):
        return value
    try:
        return json.loads(value)
    except Exception:
        return default


DISCLAIMER = (
    "Автоматичний аналіз Solomon. "
    "Підлягає перевірці юристом. Не є юридичною консультацією."
)

SEV_COLORS = {
    "critical": RGBColor(0xC0, 0x39, 0x2B),
    "high": RGBColor(0xE0, 0x6C, 0x00),
    "medium": RGBColor(0xB8, 0x89, 0x2A),
    "low": RGBColor(0x3A, 0x4D, 0x6E),
}

SEV_LABELS = {
    "critical": "КРИТИЧНИЙ",
    "high": "ВИСОКИЙ",
    "medium": "СЕРЕДНІЙ",
    "low": "НИЗЬКИЙ",
}

# Matches both Ukrainian and English AI-disclaimer tags added by the ALT prompt
_AI_TAG_RE = re.compile(
    r"\[AI\s+(?:suggestion|пропозиція)[^\]]*\]"
    r"|\(AI\s+(?:suggestion|пропозиція)[^\)]*\)",
    re.IGNORECASE,
)


def _protocol_clean(text: str, max_chars: int = 800) -> str:
    """Strip internal AI disclaimer tag and trim to 2 paragraphs / max_chars."""
    if not text:
        return ""
    text = _AI_TAG_RE.sub("", text).strip()
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    if len(paras) > 2:
        text = "\n\n".join(paras[:2])
    if len(text) > max_chars:
        text = text[:max_chars].rstrip() + "…"
    return text


def _strip_url_suffixes(url: str) -> str:
    """Normalize a zakon.rada.gov.ua URL — strip /print1, /edYYYYMMDD, #fragment."""
    if not url:
        return ""
    url = re.sub(r"/print1.*$", "", url)
    url = re.sub(r"/ed\d{8}.*$", "", url)
    url = url.split("#")[0]
    return url.rstrip("/")


def _kb_edition_map() -> dict:
    """
    Return {normalized_canonical_url: 'DD.MM.YYYY'} for active KB sources.
    Cached via get_active_kb_sources() (5-min TTL). Silently returns {} on error.
    """
    try:
        from .kb_sources import get_active_kb_sources
        sources = get_active_kb_sources()
    except Exception:
        return {}
    result = {}
    for s in sources:
        url = s.get("canonical_url", "")
        ed = s.get("current_edition_date")
        if not url or not ed:
            continue
        if hasattr(ed, "strftime"):
            ed_str = ed.strftime("%d.%m.%Y")
        else:
            try:
                y, m, d = str(ed)[:10].split("-")
                ed_str = f"{d}.{m}.{y}"
            except Exception:
                ed_str = str(ed)[:10]
        result[_strip_url_suffixes(url)] = ed_str
    return result


def _format_citations_docx(cits_raw) -> str:
    """
    Format legal_citations JSONB → compact text for DOCX.
    Each citation rendered as 'Ст. X (ред. від DD.MM.YYYY)' when edition is known.
    """
    if not cits_raw:
        return ""
    if isinstance(cits_raw, str):
        try:
            cits = json.loads(cits_raw)
        except Exception:
            return ""
    else:
        cits = cits_raw
    if not isinstance(cits, list) or not cits:
        return ""
    edition_map = _kb_edition_map()
    parts = []
    for c in cits:
        ref = c.get("article_ref") or c.get("source_title") or ""
        if not ref:
            continue
        url = c.get("official_url", "")
        ed = edition_map.get(_strip_url_suffixes(url))
        parts.append(f"{ref} (ред. від {ed})" if ed else ref)
    return "; ".join(parts)


def _set_cell_width(cell, width_cm: float):
    """Set explicit width on a table cell (twips = cm × 567)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing tcW element
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _style_header_cell(cell, text: str):
    """Bold, dark-blue text header cell."""
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0x0D, 0x15, 0x28)


def _set_table_fixed_layout(tbl, col_widths_dxa: list):
    """
    Force fixed table layout — column widths become authoritative,
    Word will not auto-expand columns to fit content.
    CT_Tbl has no get_or_add_tblPr(); must use find() + manual insert.
    """
    tbl_el = tbl._tbl

    # 1. Find or create <w:tblPr> as the FIRST child of <w:tbl>
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)

    # 2. Remove any existing <w:tblLayout>, then add fixed layout (idempotent)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # 3. Find or create <w:tblGrid> right after <w:tblPr>; rebuild gridCols
    tblGrid = tbl_el.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tblPr_index = list(tbl_el).index(tblPr)
        tbl_el.insert(tblPr_index + 1, tblGrid)
    else:
        for old in tblGrid.findall(qn("w:gridCol")):
            tblGrid.remove(old)

    # 4. One <w:gridCol w:w="…"/> per column
    for dxa in col_widths_dxa:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(dxa))
        tblGrid.append(gridCol)


def _set_cell_width_dxa(cell, dxa: int):
    """Set explicit cell width in DXA twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _style_col0_cell(cell, text: str, bold: bool = False):
    """
    Style a column-0 (№) cell: 454 DXA wide, narrow margins (40 DXA left/right),
    text centred horizontally and vertically.
    """
    _set_cell_width_dxa(cell, 454)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Narrow margins — default 120 DXA each side was eating half the column
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", 60), ("bottom", 60), ("left", 40), ("right", 40)]:
        mar_el = OxmlElement(f"w:{side}")
        mar_el.set(qn("w:w"), str(val))
        mar_el.set(qn("w:type"), "dxa")
        tcMar.append(mar_el)
    tcPr.append(tcMar)

    # Vertical centre
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)

    # Clear and write centred paragraph
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    if bold:
        run.bold = True
        run.font.color.rgb = RGBColor(0x0D, 0x15, 0x28)


# ─── §9.1 Risk note DOCX ─────────────────────────────────────────────────────

def build_risk_note_docx(
    engagement_name: str,
    counterparty: str,
    documents: list[dict],
    findings: list[dict],
) -> bytes:
    """
    Produce risk note DOCX: informal bullet list grouped by document.
    Grouped by document → category within document.
    """
    doc = Document()
    _set_margins(doc)

    h = doc.add_heading(f"Ризики — {counterparty}", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0D, 0x15, 0x28)
    doc.add_paragraph(f"Справа: {engagement_name}")
    doc.add_paragraph()

    # Group findings by document then category
    by_doc: dict[int, list] = {}
    for f in findings:
        by_doc.setdefault(f["document_id"], []).append(f)

    doc_map = {d["id"]: d for d in documents}

    for doc_id, doc_findings in by_doc.items():
        doc_info = doc_map.get(doc_id, {})
        doc.add_heading(doc_info.get("original_filename", f"Документ #{doc_id}"), level=2)

        by_cat: dict[str, list] = {}
        for f in doc_findings:
            by_cat.setdefault(f["category"], []).append(f)

        for cat, cat_findings in by_cat.items():
            doc.add_heading(_cat_label(cat), level=3)
            for f in cat_findings:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(f"[{f['clause_ref']}] ")
                r.bold = True
                sev_color = SEV_COLORS.get(f["severity"], RGBColor(0x3A, 0x4D, 0x6E))
                r2 = p.add_run(f"[{SEV_LABELS.get(f['severity'], f['severity'].upper())}] ")
                r2.font.color.rgb = sev_color
                r2.bold = True
                r3 = p.add_run(f"[{_cat_label(cat)}] ")
                r3.font.color.rgb = RGBColor(0x7A, 0x8F, 0xA8)
                p.add_run(f.get("short_note", ""))
                if f.get("monetary_exposure_uah"):
                    p.add_run(f" (≈{f['monetary_exposure_uah']:,.0f} грн)")
                if f.get("proposed_alternative"):
                    alt_p = doc.add_paragraph(style="List Bullet 2")
                    r_ai = alt_p.add_run("💡 AI пропозиція — потребує перевірки юриста: ")
                    r_ai.italic = True
                    r_ai.font.color.rgb = RGBColor(0x2E, 0x42, 0x70)
                    alt_p.add_run(f["proposed_alternative"])
                    cit_text = _format_citations_docx(f.get("legal_citations"))
                    if cit_text:
                        alt_p.add_run(f" [{cit_text}]").italic = True

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_r = footer_p.add_run(DISCLAIMER)
    footer_r.italic = True
    footer_r.font.color.rgb = RGBColor(0x7A, 0x8F, 0xA8)
    footer_r.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── §9.3 Protocol DOCX ──────────────────────────────────────────────────────

def build_protocol_docx(
    engagement_name: str,
    counterparty: str,
    findings: list[dict],
    document_filename: str = "",
    avtd_role: Optional[str] = None,
) -> bytes:
    """
    4-column counterparty-facing negotiation table.

    Columns: № | Редакція {counterparty} | Редакція {AVTD} | Узгоджена редакція

    - avtd_role: 'supplier' or 'buyer'. Required — raises ValueError if None.
    - Column 2 header: dynamic (counterparty's role label).
    - Column 3 header: dynamic (AVTD's role label).
    - AI disclaimer tag stripped from BOTH counterparty and AVTD columns.
    - Clause ref (e.g. "п.13.5") moved to bold prefix at top of column 2.
    - Правова підстава column DROPPED — stays in solcon_findings.legal_citations
      and is rendered in the правовий висновок DOCX only.
    - Column widths: №=0.8cm, counterparty≈6.0cm, AVTD≈6.0cm, agreed≈5.2cm
      (total ≈18.0cm, fits A4 with 1.2cm/1.0cm margins).
    """
    if not avtd_role:
        raise ValueError(
            "Set AVTD role on engagement before exporting protocol. "
            "Use PATCH /api/contracts/engagements/{eid} with {\"avtd_role\": \"supplier\"|\"buyer\"}."
        )
    if avtd_role not in ("supplier", "buyer"):
        raise ValueError(f"Invalid avtd_role '{avtd_role}'. Must be 'supplier' or 'buyer'.")

    # Dynamic column headers based on AVTD role
    if avtd_role == "supplier":
        counterparty_col_header = "Редакція Покупця"
        avtd_col_header = "Редакція Постачальника"
        role_label_ua = "Постачальником"
    else:
        counterparty_col_header = "Редакція Постачальника"
        avtd_col_header = "Редакція Покупця"
        role_label_ua = "Покупцем"

    doc = Document()
    _set_margins(doc)

    h = doc.add_heading("ПРОТОКОЛ РОЗБІЖНОСТЕЙ", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0D, 0x15, 0x28)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_lines = [f"до Договору поставки з {counterparty}"]
    if document_filename:
        subtitle_lines.append(f"документ: {document_filename}")
    subtitle_lines += [
        f"Справа: {engagement_name}",
        f"AVTD виступає {role_label_ua}.",
    ]
    doc.add_paragraph("\n".join(subtitle_lines))
    doc.add_paragraph()

    # 4-column table: №  | counterparty edition | AVTD edition | agreed
    # 454 + 3200 + 3200 + 2506 = 9360 DXA ≈ 16.5cm (fits A4 with 1.2cm/1.0cm margins)
    COL_WIDTHS_DXA = [454, 3200, 3200, 2506]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _set_table_fixed_layout(tbl, COL_WIDTHS_DXA)  # forces Word to honour narrow col 0

    hdr = tbl.rows[0].cells
    headers = ["№", counterparty_col_header, avtd_col_header, "Узгоджена редакція"]

    _style_col0_cell(hdr[0], "№", bold=True)       # narrow margins + centred
    for i in range(1, 4):
        _style_header_cell(hdr[i], headers[i])
        _set_cell_width_dxa(hdr[i], COL_WIDTHS_DXA[i])

    for ordinal, f in enumerate(findings, start=1):
        row = tbl.add_row().cells

        # Col 0 — ordinal №: narrow, centred, vertically centred
        _style_col0_cell(row[0], str(ordinal))

        # Col 1 — counterparty's verbatim text, clause ref as bold prefix
        clause_ref = (f.get("clause_ref") or "").strip()
        clause_text = _protocol_clean(f.get("clause_text") or "")
        col1_para = row[1].paragraphs[0]
        if clause_ref:
            ref_run = col1_para.add_run(f"{clause_ref}\n")
            ref_run.bold = True
        col1_para.add_run(clause_text)
        _set_cell_width_dxa(row[1], COL_WIDTHS_DXA[1])

        # Col 2 — AVTD's alternative (AI tag stripped, capped at 800 chars)
        avtd_raw = f.get("proposed_alternative") or f.get("short_note", "")
        row[2].text = _protocol_clean(avtd_raw)
        _set_cell_width_dxa(row[2], COL_WIDTHS_DXA[2])

        # Col 3 — Узгоджена редакція: blank for manual completion
        row[3].text = ""
        _set_cell_width_dxa(row[3], COL_WIDTHS_DXA[3])

    doc.add_paragraph()
    doc.add_paragraph(
        "Цей протокол складено у двох примірниках, по одному для кожної Сторони."
    )
    footer_p = doc.add_paragraph()
    footer_r = footer_p.add_run(DISCLAIMER)
    footer_r.italic = True
    footer_r.font.color.rgb = RGBColor(0x7A, 0x8F, 0xA8)
    footer_r.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── §9.2 Legal opinion DOCX (from markdown) ─────────────────────────────────

def build_opinion_docx(
    markdown_text: str,
    engagement_name: str,
    findings: Optional[list] = None,
) -> bytes:
    """
    Convert markdown legal opinion to DOCX.
    If `findings` are provided, appends a 'Правова база' table listing every
    cited KB source with its current edition date and last-verified timestamp.
    """
    doc = Document()
    _set_margins(doc)

    for line in markdown_text.split("\n"):
        if line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x0D, 0x15, 0x28)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _add_md_run(p, line)

    # ── Правова база appendix ────────────────────────────────────────────────
    if findings:
        _append_kb_sources_table(doc, findings)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _append_kb_sources_table(doc: "Document", findings: list) -> None:
    """
    Append a 'Правова база' section to the legal opinion DOCX.
    Collects all unique cited URLs from findings, matches them against
    the active KB registry, and renders a 3-column table:
    Закон | Редакція | Перевірено
    """
    try:
        from .kb_sources import get_active_kb_sources
        sources = get_active_kb_sources()
    except Exception:
        return

    # Build lookup: normalized_url → source row
    src_by_url: dict = {_strip_url_suffixes(s["canonical_url"]): s for s in sources if s.get("canonical_url")}

    # Collect unique cited URLs from all findings
    seen_urls: set = set()
    cited_sources: list = []
    for f in findings:
        cits = _safe_jsonb(f.get("legal_citations"))
        for c in (cits or []):
            norm = _strip_url_suffixes(c.get("official_url", ""))
            if norm and norm not in seen_urls:
                seen_urls.add(norm)
                src = src_by_url.get(norm)
                if src:
                    cited_sources.append(src)

    if not cited_sources:
        return

    doc.add_paragraph()
    h = doc.add_heading("Правова база", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x0D, 0x15, 0x28)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _set_table_fixed_layout(tbl, [5670, 1701, 1701])  # 10cm + 3cm + 3cm = 16cm

    hdr_cells = tbl.rows[0].cells
    for cell, text in zip(hdr_cells, ["Закон / джерело", "Редакція", "Перевірено"]):
        _style_header_cell(cell, text)

    for src in cited_sources:
        row = tbl.add_row().cells
        row[0].text = src.get("law_name", src.get("law_code", ""))

        ed = src.get("current_edition_date")
        if ed:
            if hasattr(ed, "strftime"):
                ed_str = ed.strftime("%d.%m.%Y")
            else:
                try:
                    y, m, d = str(ed)[:10].split("-")
                    ed_str = f"{d}.{m}.{y}"
                except Exception:
                    ed_str = str(ed)[:10]
        else:
            ed_str = "—"
        row[1].text = ed_str

        ver = src.get("last_verified_at")
        if ver:
            if hasattr(ver, "strftime"):
                ver_str = ver.strftime("%d.%m.%Y")
            else:
                try:
                    ver_str = str(ver)[:10].replace("-", ".")
                    y, m, d = str(ver)[:10].split("-")
                    ver_str = f"{d}.{m}.{y}"
                except Exception:
                    ver_str = str(ver)[:10]
        else:
            ver_str = "—"
        row[2].text = ver_str


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _add_md_run(p, text: str):
    """Minimal bold/italic markdown parsing for inline text."""
    import re
    segments = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for seg in segments:
        if seg.startswith("**") and seg.endswith("**"):
            r = p.add_run(seg[2:-2])
            r.bold = True
        elif seg.startswith("*") and seg.endswith("*"):
            r = p.add_run(seg[1:-1])
            r.italic = True
        else:
            p.add_run(seg)


def _set_margins(doc: Document):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)


def _cat_label(cat: str) -> str:
    labels = {
        "penalty": "Штрафні санкції",
        "payment_terms": "Умови оплати",
        "liability_shift": "Перенесення відповідальності",
        "ip_rights": "Права інтелектуальної власності",
        "force_majeure": "Форс-мажор",
        "termination": "Розірвання договору",
        "returns_refusal": "Повернення / відмова товару",
        "audit_rights": "Право аудиту",
        "set_off": "Залік вимог",
        "tax_invoicing": "Податкові накладні",
        "quality_acceptance": "Приймання за якістю",
        "delivery_terms": "Умови поставки",
        "other": "Інше",
    }
    return labels.get(cat, cat)
