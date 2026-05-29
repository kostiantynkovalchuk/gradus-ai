"""
Bundle ingestion pipeline for Solomon Contracts.
Handles: file saving, text extraction, document classification, clause parsing.
"""
import io
import json
import logging
import os
import re
import subprocess
import tempfile
import time
import zipfile
from pathlib import Path
from typing import Optional

import anthropic
from services.ai_models import HAIKU

logger = logging.getLogger(__name__)

_DEFAULT_UPLOAD_DIR = str(Path(__file__).parent.parent / "solomon_uploads")
UPLOAD_DIR = Path(os.getenv("SOLCON_UPLOAD_DIR", _DEFAULT_UPLOAD_DIR))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

CLAUSE_RE = re.compile(
    r"(?m)^(?P<ref>"
    r"п\.\s*\d+(?:\.\d+){1,3}\.?"
    r"|\d+(?:\.\d+){1,3}\.?"
    r"|Розділ\s+\d+"
    r"|Додаток\s*№?\s*\d+"
    r")\s*"
)

DOC_TYPE_HINTS = [
    (re.compile(r"риз(ики|ики)", re.I), "risks_note"),
    (re.compile(r"висновок", re.I), "legal_opinion"),
    (re.compile(r"протокол", re.I), "protocol_draft"),
    (re.compile(r"кодекс|ethics|code", re.I), "commercial_code"),
    (re.compile(r"специфікац|специф", re.I), "specification"),
    (re.compile(r"прайс|price[._\s-]?list", re.I), "price_list"),
    (re.compile(r"додаткова.угода|ду\b|допка|edi", re.I), "additional_agreement"),
]

VALID_DOC_TYPES = {
    "main_contract", "additional_agreement", "commercial_code",
    "specification", "price_list", "schedule", "risks_note", "legal_opinion",
    "protocol_draft", "protocol_returned", "protocol_agreed", "other",
}


# ─── Text extraction ─────────────────────────────────────────────────────────

def extract_text_from_docx(path: Path) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paras.append(row_text)
    return "\n".join(paras)


def extract_text_from_doc(path: Path) -> str:
    """Try antiword → LibreOffice → python-docx fallback chain."""
    # 1. antiword (fastest, no headless overhead)
    try:
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout:
            return result.stdout.decode("utf-8", errors="replace")
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    # 2. LibreOffice headless conversion
    try:
        with tempfile.TemporaryDirectory() as tmp:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmp, str(path)],
                capture_output=True, timeout=30,
            )
            converted = list(Path(tmp).glob("*.docx"))
            if result.returncode == 0 and converted:
                return extract_text_from_docx(converted[0])
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    # 3. python-docx last-ditch (works on some .doc that are actually .docx-zipped)
    try:
        return extract_text_from_docx(path)
    except Exception:
        return ""


_OCR_THRESHOLD = 500  # chars — below this pdftotext result triggers OCR fallback


_OCR_DPI = 150          # 150 dpi: ~44% memory vs 300 dpi; ~2× faster per page; legible for dense legal text
_TESS_CONFIG = "--psm 6"  # PSM 6 = assume single uniform text block; 30-50% faster than default PSM 3 on dense pages
_OCR_MAX_PAGES = 40    # hard cap — override with SOLCON_OCR_MAX_PAGES env var


def _pdf_page_count(path: Path) -> int:
    """Return number of pages using pdfinfo (poppler). Returns 0 on failure."""
    try:
        result = subprocess.run(
            ["pdfinfo", str(path)], capture_output=True, timeout=10,
        )
        if result.returncode == 0:
            for line in result.stdout.decode("utf-8", errors="replace").splitlines():
                if line.startswith("Pages:"):
                    return int(line.split(":", 1)[1].strip())
    except Exception as e:
        logger.debug("[SolCon] pdfinfo failed for %s: %s", path.name, e)
    return 0


def _ocr_pdf(path: Path) -> tuple[str, str]:
    """
    Stream-OCR: rasterize and OCR one page at a time to keep peak RSS low.
    Each page image is freed immediately after OCR — peak memory ≈ 1 page
    at _OCR_DPI dpi (~10–20 MB for A4) instead of all pages at once.

    Returns (full_text, method_label).
    """
    try:
        import pdf2image
        import pytesseract

        max_pages = int(os.getenv("SOLCON_OCR_MAX_PAGES", _OCR_MAX_PAGES))
        page_count = _pdf_page_count(path)
        if page_count == 0:
            logger.warning("[SolCon] pdfinfo returned 0 pages for %s — attempting blind OCR", path.name)
            page_count = max_pages   # will break early when convert_from_path returns []

        pages_to_ocr = min(page_count, max_pages)
        if page_count > max_pages:
            logger.warning(
                "[SolCon] %s has %d pages — capping OCR at %d (SOLCON_OCR_MAX_PAGES)",
                path.name, page_count, max_pages,
            )

        t0 = time.time()
        pages_text: list[str] = []

        for page_num in range(1, pages_to_ocr + 1):
            # Rasterize exactly one page — only that one image lives in RAM
            images = pdf2image.convert_from_path(
                str(path), dpi=_OCR_DPI,
                first_page=page_num, last_page=page_num,
            )
            if not images:
                break  # no more pages (blind-OCR mode hit the real end)

            img = images[0]
            del images  # release list; img still referenced

            pt0 = time.time()
            text = pytesseract.image_to_string(img, lang="ukr+rus")
            elapsed = time.time() - pt0

            del img   # free the rasterized image buffer immediately

            logger.info(
                "[SolCon] OCR page %d/%d: %d chars in %.1fs — %s",
                page_num, pages_to_ocr, len(text), elapsed, path.name,
            )
            if text.strip():
                pages_text.append(f"[Стор. {page_num}]\n{text}")

        full_text = "\n\n".join(pages_text)
        total = time.time() - t0
        logger.info(
            "[SolCon] OCR done: %d chars, %d/%d pages, %.1fs total — %s",
            len(full_text), len(pages_text), pages_to_ocr, total, path.name,
        )
        return full_text, "ocr:tesseract:ukr+rus"

    except Exception as e:
        logger.warning("[SolCon] OCR failed for %s: %s", path.name, e)
        return "", "ocr:failed"


def extract_text_from_pdf(path: Path) -> tuple[str, str]:
    """
    Returns (text, extraction_method).
    Chain (fast, synchronous — no OCR here):
      1. pdftotext → if ≥500 chars: done
      2. pdfplumber → if ≥500 chars: done
      3. pdfminer   → if ≥500 chars: done
      4. Scanned PDF detected → return ("", "ocr:pending")
         Caller (ingest_file / re-extract) schedules run_background_ocr().
    """
    # 1. pdftotext (poppler) — most reliable for text-based Ukrainian PDFs
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(path), "-"],
            capture_output=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout.decode("utf-8", errors="replace")
            logger.info("[SolCon] pdftotext: %d chars from %s", len(text), path.name)
            if len(text) >= _OCR_THRESHOLD:
                return text, "pdftotext"
            logger.info(
                "[SolCon] pdftotext: only %d chars (<500) for %s — likely scanned",
                len(text), path.name,
            )
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        logger.debug("[SolCon] pdftotext unavailable or timed out: %s", e)

    # 2. pdfplumber (catches some PDFs with unusual encoding that pdftotext misses)
    try:
        import pdfplumber
        pages: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text() or ""
                if t:
                    pages.append(f"[Стор. {i + 1}]\n{t}")
        text = "\n\n".join(pages)
        if len(text) >= _OCR_THRESHOLD:
            logger.info("[SolCon] pdfplumber: %d chars from %s", len(text), path.name)
            return text, "pdfplumber"
    except Exception as e:
        logger.warning("[SolCon] pdfplumber failed for %s: %s", path.name, e)

    # 3. pdfminer last native-text attempt
    try:
        from pdfminer.high_level import extract_text as _pdfminer
        text = _pdfminer(str(path)) or ""
        if len(text) >= _OCR_THRESHOLD:
            logger.info("[SolCon] pdfminer: %d chars from %s", len(text), path.name)
            return text, "pdfminer"
    except Exception as e:
        logger.warning("[SolCon] pdfminer failed for %s: %s", path.name, e)

    # 4. All native extractors exhausted → PDF is scanned; OCR runs in background
    logger.info("[SolCon] '%s' is scanned — signalling ocr:pending for async OCR", path.name)
    return "", "ocr:pending"


def run_background_ocr(doc_id: int, path: Path, delete_when_done: bool = False) -> None:
    """
    Stream-OCR a scanned PDF page-by-page in a daemon thread.

    Memory: peak ≈ 1 page × _OCR_DPI dpi (~5–10 MB) — only one PIL image
    lives in RAM at a time.  Progress written to DB after every page so the
    UI can poll ocr_current_page/ocr_total_pages.

    On success:  raw_text, clauses, extraction_method, ocr_status='done'
    On failure:  ocr_status='failed'

    delete_when_done: if True, the thread deletes the file at `path` after OCR
    completes or fails (used when path is a NamedTemporaryFile from re-extract).
    """
    import threading
    import json as _json
    from . import db as _db

    def _run():
        try:
            import pdf2image
            import pytesseract

            max_pages = int(os.getenv("SOLCON_OCR_MAX_PAGES", _OCR_MAX_PAGES))
            page_count = _pdf_page_count(path)
            if page_count == 0:
                logger.warning(
                    "[SolCon] BG OCR: pdfinfo returned 0 for doc_id=%d — trying up to %d pages",
                    doc_id, max_pages,
                )
                page_count = max_pages

            pages_to_ocr = min(page_count, max_pages)
            if page_count > max_pages:
                logger.warning(
                    "[SolCon] BG OCR: doc_id=%d has %d pages — capping at %d (SOLCON_OCR_MAX_PAGES)",
                    doc_id, page_count, max_pages,
                )

            _db.execute(
                "UPDATE solcon_documents SET ocr_status='running', ocr_total_pages=%s, ocr_current_page=0 WHERE id=%s",
                (pages_to_ocr, doc_id),
            )
            logger.info("[SolCon] BG OCR started: doc_id=%d, %d pages, %d dpi, config=%r",
                        doc_id, pages_to_ocr, _OCR_DPI, _TESS_CONFIG)

            t0 = time.time()
            pages_text: list[str] = []

            for page_num in range(1, pages_to_ocr + 1):
                images = pdf2image.convert_from_path(
                    str(path), dpi=_OCR_DPI,
                    first_page=page_num, last_page=page_num,
                )
                if not images:
                    logger.warning("[SolCon] BG OCR: no image for page %d of doc_id=%d — stopping", page_num, doc_id)
                    break

                img = images[0]
                del images

                pt0 = time.time()
                text = pytesseract.image_to_string(img, lang="ukr+rus", config=_TESS_CONFIG)
                elapsed = time.time() - pt0
                del img  # free rasterized buffer immediately

                logger.info("[SolCon] BG OCR doc_id=%d page %d/%d: %d chars in %.1fs",
                            doc_id, page_num, pages_to_ocr, len(text), elapsed)

                if text.strip():
                    pages_text.append(f"[Стор. {page_num}]\n{text}")

                _db.execute(
                    "UPDATE solcon_documents SET ocr_current_page=%s WHERE id=%s",
                    (page_num, doc_id),
                )

            full_text = "\n\n".join(pages_text)
            total = time.time() - t0

            clauses = parse_clauses(full_text)
            if not clauses:
                clauses = scan_all_refs(full_text)

            _db.execute(
                """UPDATE solcon_documents
                   SET raw_text=%s, clauses=%s::jsonb,
                       extraction_method='ocr:tesseract:ukr+rus',
                       ocr_status='done', ocr_current_page=%s,
                       analyzed_at=NULL, updated_at=NOW()
                   WHERE id=%s""",
                (full_text, _json.dumps(clauses), pages_to_ocr, doc_id),
            )
            logger.info("[SolCon] BG OCR done: doc_id=%d, %d chars, %.1fs total",
                        doc_id, len(full_text), total)

        except Exception as exc:
            logger.exception("[SolCon] BG OCR FAILED for doc_id=%d: %s", doc_id, exc)
            try:
                _db.execute(
                    "UPDATE solcon_documents SET ocr_status='failed' WHERE id=%s",
                    (doc_id,),
                )
            except Exception:
                pass
        finally:
            if delete_when_done:
                try:
                    path.unlink(missing_ok=True)
                except Exception:
                    pass

    threading.Thread(target=_run, daemon=True, name=f"ocr-doc-{doc_id}").start()


def extract_text_from_xlsx(path: Path) -> str:
    """Handles modern .xlsx files via openpyxl."""
    from openpyxl import load_workbook
    rows = []
    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        for ws in wb.worksheets:
            rows.append(f"[Аркуш: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append(" | ".join(cells))
        wb.close()
        logger.info(f"[SolCon] openpyxl: {len(rows)} rows from {path.name}")
    except Exception as e:
        logger.warning(f"[SolCon] openpyxl extraction failed for {path.name}: {e}")
    return "\n".join(rows)


def extract_text_from_xls(path: Path) -> str:
    """Handles legacy .xls files via xlrd (openpyxl cannot read old XLS format)."""
    try:
        import xlrd
        wb = xlrd.open_workbook(str(path))
        rows = []
        for sheet in wb.sheets():
            rows.append(f"[Аркуш: {sheet.name}]")
            for rowx in range(sheet.nrows):
                cells = []
                for c in range(sheet.ncols):
                    ct = sheet.cell_type(rowx, c)
                    if ct not in (xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK):
                        cells.append(str(sheet.cell_value(rowx, c)))
                if cells:
                    rows.append(" | ".join(cells))
        text = "\n".join(rows)
        logger.info(f"[SolCon] xlrd: {len(rows)} rows / {len(text)} chars from {path.name}")
        return text
    except Exception as e:
        logger.warning(f"[SolCon] xlrd extraction failed for {path.name}: {e}")
        return ""


def extract_text(path: Path, mime_type: str = "") -> tuple[str, Optional[str]]:
    """
    Returns (raw_text, extraction_method).
    extraction_method is None for non-PDF types; for PDFs it reflects which
    extractor succeeded: 'pdftotext', 'ocr:tesseract:ukr+rus', 'pdfplumber',
    'pdfminer', or 'failed'.
    """
    suffix = path.suffix.lower()
    if suffix == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(path), None
    if suffix == ".doc" or mime_type == "application/msword":
        return extract_text_from_doc(path), None
    if suffix == ".pdf" or mime_type == "application/pdf":
        return extract_text_from_pdf(path)   # already returns (text, method)
    if suffix == ".xls" or mime_type == "application/vnd.ms-excel":
        return extract_text_from_xls(path), None    # legacy format — must use xlrd
    if suffix == ".xlsx" or "spreadsheet" in mime_type:
        return extract_text_from_xlsx(path), None   # modern format — openpyxl
    logger.warning(f"[SolCon] Unknown file type {suffix} — skipping extraction")
    return "", None


# ─── Document classification ─────────────────────────────────────────────────

def _classify_by_filename(filename: str) -> Optional[str]:
    name = filename.lower()
    ext = Path(filename).suffix.lower()
    is_spreadsheet = ext in (".xlsx", ".xls", ".ods", ".csv")
    is_pdf = ext == ".pdf"

    # --- Extension-specific patterns (checked before generic DOC_TYPE_HINTS) ---
    if is_spreadsheet:
        if re.search(r"специф", name):   return "specification"
        if re.search(r"прайс|price",name): return "price_list"
        if re.search(r"графік|график", name): return "schedule"
        if re.search(r"накладна", name):  return "other"   # sample delivery doc

    if is_pdf and re.search(r"додаток|додатки", name):
        return "additional_agreement"

    # --- Generic DOC_TYPE_HINTS (order matters) ---
    for pattern, doc_type in DOC_TYPE_HINTS:
        if pattern.search(name):
            return doc_type

    # --- Broader main-contract patterns ---
    if re.search(r"договір|договор", name):
        return "main_contract"
    if re.search(r"т\.ф\.\s*\d+", name):          # т.ф. 188 retail form
        return "main_contract"

    # --- Catch-all other ---
    if re.search(r"пакувальний|packing", name):
        return "other"

    return None


def _classify_via_llm(filename: str, text_snippet: str) -> str:
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    t0 = time.time()
    msg = client.messages.create(
        model=HAIKU,
        max_tokens=64,
        system=(
            "You classify Ukrainian legal documents. "
            "Reply with exactly ONE of: main_contract, additional_agreement, "
            "commercial_code, specification, price_list, risks_note, legal_opinion, "
            "protocol_draft, protocol_returned, protocol_agreed, other. "
            "No explanation."
        ),
        messages=[{
            "role": "user",
            "content": f"Filename: {filename}\n\nFirst 500 chars:\n{text_snippet[:500]}",
        }],
    )
    duration_ms = int((time.time() - t0) * 1000)
    result = msg.content[0].text.strip().lower()
    from . import db as solcon_db
    solcon_db.log_llm_call(
        None, None, "classify", HAIKU,
        msg.usage.input_tokens, msg.usage.output_tokens, duration_ms,
    )
    return result if result in VALID_DOC_TYPES else "other"


def classify_document(filename: str, raw_text: str) -> str:
    logger.info(
        "[SolCon] Classify '%s': extracted %d chars. Preview: %r",
        filename, len(raw_text), raw_text[:500],
    )
    hint = _classify_by_filename(filename)
    if hint:
        logger.info("[SolCon] Filename heuristic '%s' → '%s'", filename, hint)
        return hint
    if not raw_text.strip():
        logger.warning("[SolCon] '%s': empty text after extraction → 'other'", filename)
        return "other"
    return _classify_via_llm(filename, raw_text)


# ─── Clause parsing ──────────────────────────────────────────────────────────

CLAUSE_ANYWHERE_RE = re.compile(
    r"(?:п\.\s*)?(?<!\d)(\d{1,2}\.\d{1,2}(?:\.\d{1,2}){0,2})(?!\d)"
)

def scan_all_refs(raw_text: str, appendix_prefix: str = "") -> list[dict]:
    """
    Find all X.Y[.Z] clause refs mentioned ANYWHERE in the text (cross-refs, etc.)
    Used to supplement line-start parsing for table-formatted documents.
    Returns list of {ref, text, parent_ref} with empty text (ref-only entries).
    Filters out obvious dates (dd.mm.yyyy pattern).
    """
    DATE_RE = re.compile(r"\d{2}\.\d{2}\.\d{4}")
    date_spans = {m.span() for m in DATE_RE.finditer(raw_text)}

    seen = set()
    clauses = []
    for m in CLAUSE_ANYWHERE_RE.finditer(raw_text):
        if m.span() in date_spans:
            continue
        ref = m.group(0).strip()
        norm = re.sub(r"п\.\s*", "п.", ref).rstrip(".")
        if norm in seen:
            continue
        seen.add(norm)
        clauses.append({
            "ref": appendix_prefix + norm,
            "text": "",
            "parent_ref": _parent_ref(norm),
        })
    return clauses


def parse_clauses(raw_text: str, appendix_prefix: str = "") -> list[dict]:
    """
    Extract clause refs from text using Ukrainian contract numbering conventions.
    Returns list of {ref, text, parent_ref}.
    """
    clauses = []
    lines = raw_text.split("\n")
    current_ref = None
    current_lines: list[str] = []

    def _flush():
        if current_ref:
            clauses.append({
                "ref": (appendix_prefix + current_ref).strip(),
                "text": " ".join(current_lines).strip(),
                "parent_ref": _parent_ref(current_ref),
            })

    for line in lines:
        m = CLAUSE_RE.match(line)
        if m:
            _flush()
            current_ref = m.group("ref").strip()
            current_lines = [line[m.end():].strip()]
        else:
            if current_ref:
                current_lines.append(line.strip())

    _flush()
    return clauses


def _parent_ref(ref: str) -> Optional[str]:
    """п.5.5.1 → п.5.5, п.5.5 → п.5, п.5 → None"""
    m = re.match(r"(п\.\s*)?(\d+(?:\.\d+)*)", ref)
    if not m:
        return None
    parts = m.group(2).split(".")
    if len(parts) <= 1:
        return None
    prefix = "п." if ref.startswith("п") else ""
    return prefix + ".".join(parts[:-1])


def normalize_clause_ref(ref: str) -> str:
    """Normalize 'п. 5.5' → 'п.5.5'"""
    ref = re.sub(r"п\.\s+", "п.", ref)
    return ref.strip()


# ─── Bundle ingestion ────────────────────────────────────────────────────────

def save_upload(filename: str, data: bytes, engagement_id: int) -> tuple:
    eng_dir = UPLOAD_DIR / str(engagement_id)
    eng_dir.mkdir(parents=True, exist_ok=True)
    dest = eng_dir / filename
    dest.write_bytes(data)
    return dest, data


def process_zip(data: bytes, engagement_id: int) -> list[dict]:
    """Extract zip, return list of {filename, path, bytes, size}."""
    files = []
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for member in zf.infolist():
            if member.is_dir():
                continue
            name = Path(member.filename).name
            suffix = Path(name).suffix.lower()
            if suffix not in {".doc", ".docx", ".pdf", ".xls", ".xlsx"}:
                continue
            content = zf.read(member.filename)
            path, file_bytes = save_upload(name, content, engagement_id)
            files.append({"filename": name, "path": path, "bytes": file_bytes, "size": len(content)})
    return files


def ingest_file(
    filename: str, data: bytes, engagement_id: int,
    appendix_prefix: str = "",
) -> dict:
    """
    Save file, extract text, classify, parse clauses.
    Returns dict ready for DB insert into solcon_documents.

    When a scanned PDF is detected, extraction_method='ocr:pending' is returned.
    The caller (router upload endpoint) must then call run_background_ocr(doc_id, path).
    """
    path, file_bytes = save_upload(filename, data, engagement_id)
    suffix = Path(filename).suffix.lower()
    mime_map = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".pdf": "application/pdf",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
    }
    mime_type = mime_map.get(suffix, "application/octet-stream")
    raw_text, extraction_method = extract_text(path, mime_type)

    ocr_pending = extraction_method == "ocr:pending"

    # Classification uses filename heuristics first — works even with empty text
    doc_type = classify_document(filename, raw_text)
    # Clause parsing on empty text produces nothing useful; skip for OCR-pending docs
    clauses = [] if ocr_pending else parse_clauses(raw_text, appendix_prefix)

    return {
        "engagement_id": engagement_id,
        "document_type": doc_type,
        "original_filename": filename,
        "mime_type": mime_type,
        "storage_path": str(path),
        "file_bytes": file_bytes,
        "raw_text": raw_text,
        "clauses": json.dumps(clauses),
        # extraction_method is NULL in DB until OCR finishes and sets 'ocr:tesseract:ukr+rus'
        "extraction_method": None if ocr_pending else extraction_method,
        # ocr_status drives the progress UI; None for non-scanned docs
        "ocr_status": "pending" if ocr_pending else None,
        "_ocr_pending": ocr_pending,   # internal flag — NOT stored in DB directly
        "_path": path,                  # internal — passed to run_background_ocr
    }
