"""
Solomon Contracts — Automated Phase-Gate Eval (§13)

Usage:
    cd backend && python -m solomon_contracts.eval

Reads from training_data/levays/:
  - 1АЛ_Договір_поставки_умови_Вчасно_ЛЕВАЙС.doc   (main contract to analyze)
  - РИЗИКИ_ДО_ДОГОВОРУ_ПОСТАВКИ_ЛЕВАЙС_10_06_2025.docx  (ground-truth risk note)

Computes (§13.2):
  - Precision  ≥ 0.75 target
  - Recall     ≥ 0.70 target
  - Clause-ref accuracy = 1.0 target  (every Solomon finding cites a real clause)

Note: Grounding rate is NOT measured here — corpus is not yet seeded.
      Run again after corpus seed (#7) for grounding metrics.

Outputs:
  - Console summary
  - training_data/levays/eval_results.json
"""
import json
import logging
import os
import re
import sys
import time
from pathlib import Path
from typing import Optional

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s: %(message)s")
logger = logging.getLogger("solcon.eval")

ROOT = Path(__file__).parent.parent.parent
LEVAYS_DIR = ROOT / "training_data" / "levays"
RESULTS_FILE = LEVAYS_DIR / "eval_results.json"

EXPECTED_CONTRACT = "1АЛ_Договір_поставки_умови_Вчасно_ЛЕВАЙС.doc"
EXPECTED_RISK_NOTE = "РИЗИКИ_ДО_ДОГОВОРУ_ПОСТАВКИ_ЛЕВАЙС_10_06_2025.docx"

PRECISION_TARGET = 0.75
RECALL_TARGET = 0.70
CLAUSE_ACC_TARGET = 1.0

CLAUSE_RE = re.compile(
    r"(?P<ref>"
    r"п\.\s*\d+(?:\.\d+){1,3}\.?"
    r"|\d+(?:\.\d+){1,3}\.?"
    r"|Розділ\s+\d+"
    r"|Додаток\s*№?\s*\d+"
    r")"
)

CATEGORY_KEYWORDS = {
    "penalty":           ["штраф", "пеня", "неустойка", "санкц"],
    "payment_terms":     ["оплат", "розрахун", "відстрочк", "платіж"],
    "liability_shift":   ["відповідальн", "збитк", "ризик"],
    "force_majeure":     ["форс", "непередбачен"],
    "termination":       ["розірван", "припинен", "відмов"],
    "returns_refusal":   ["поверн", "відмов", "рекламац"],
    "audit_rights":      ["перевірк", "аудит", "контрол"],
    "set_off":           ["залік", "зустрічн"],
    "tax_invoicing":     ["податков", "накладн", "ПДВ"],
    "quality_acceptance":["якість", "якост", "прийом", "приймання"],
    "delivery_terms":    ["доставк", "поставк", "термін"],
    "ip_rights":         ["торгов", "марк", "бренд", "авторськ"],
}


# ─── Ground-truth parser ─────────────────────────────────────────────────────

def parse_risk_note(path: Path) -> list[dict]:
    """
    Parse lawyer-written risk note DOCX into structured findings.
    Strategy:
      1. If the doc has a table with clause_ref column, parse that.
      2. Otherwise, extract clause refs from paragraphs and surrounding text.
    Returns list of {clause_ref, category, raw_text}.
    """
    from docx import Document
    doc = Document(str(path))

    findings = []

    # Strategy 1: table parsing (4+ columns suggests a structured risk register)
    for table in doc.tables:
        if len(table.columns) >= 2:
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:
                    continue  # skip header
                cells = [c.text.strip() for c in row.cells]
                ref_candidates = []
                for cell in cells:
                    for m in CLAUSE_RE.finditer(cell):
                        ref_candidates.append(m.group("ref"))
                if ref_candidates:
                    raw = " ".join(cells)
                    findings.append({
                        "clause_ref": normalize_ref(ref_candidates[0]),
                        "category": _classify_text(raw),
                        "raw_text": raw[:300],
                        "source": "table",
                    })

    if findings:
        logger.info(f"Ground truth: extracted {len(findings)} findings from table(s)")
        return _deduplicate(findings)

    # Strategy 2: paragraph scanning
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        refs = list(CLAUSE_RE.finditer(para))
        if refs:
            # Collect context: this paragraph + next 2
            context = " ".join(paragraphs[i:i + 3])
            for m in refs:
                ref = m.group("ref")
                findings.append({
                    "clause_ref": normalize_ref(ref),
                    "category": _classify_text(context),
                    "raw_text": context[:300],
                    "source": "paragraph",
                })
        i += 1

    findings = _deduplicate(findings)
    logger.info(f"Ground truth: extracted {len(findings)} findings from paragraphs")
    return findings


def _classify_text(text: str) -> str:
    text_lower = text.lower()
    best = "other"
    best_score = 0
    for cat, keywords in CATEGORY_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score > best_score:
            best_score = score
            best = cat
    return best


def _deduplicate(findings: list[dict]) -> list[dict]:
    seen = set()
    out = []
    for f in findings:
        key = f["clause_ref"]
        if key not in seen:
            seen.add(key)
            out.append(f)
    return out


# ─── Solomon analyzer runner ─────────────────────────────────────────────────

def run_solomon_on_contract(contract_path: Path) -> tuple[list[dict], list[dict]]:
    """
    Extract text, parse clauses, run free-form scan.
    Returns (solomon_findings, parsed_clauses).
    No DB is required for metrics — we stub engagement/document IDs.
    LLM audit calls still log to DB if reachable (non-critical).
    """
    from .ingestion import extract_text, parse_clauses
    from .analyzer import scan_document

    logger.info(f"Extracting text from: {contract_path.name}")
    raw_text = extract_text(contract_path)
    if not raw_text.strip():
        logger.error("Contract text extraction returned empty string")
        sys.exit(1)
    logger.info(f"Extracted {len(raw_text)} characters, {len(raw_text.split())} words")

    logger.info("Parsing clause references…")
    clauses = parse_clauses(raw_text)
    logger.info(f"Parsed {len(clauses)} clauses")

    logger.info("Running Solomon free-form scan (Claude Sonnet)… this takes ~30-90s")
    t0 = time.time()
    findings, rejected_count = scan_document(
        document_id=0,
        engagement_id=0,
        raw_text=raw_text,
        clauses=clauses,
    )
    elapsed = time.time() - t0
    logger.info(
        f"Scan complete in {elapsed:.1f}s — "
        f"{len(findings)} findings accepted, {rejected_count} rejected by guardrail §10.1"
    )
    return findings, clauses


# ─── Metrics ─────────────────────────────────────────────────────────────────

def normalize_ref(ref: str) -> str:
    """Normalize clause ref for matching: lowercase, remove spaces."""
    ref = ref.strip().lower()
    ref = re.sub(r"п\.\s+", "п.", ref)
    ref = re.sub(r"\.\s+", ".", ref)
    ref = ref.rstrip(".")
    return ref


def _refs_match(ref_a: str, ref_b: str) -> bool:
    """
    Consider refs matching if:
    - Exact normalized match, OR
    - One starts with the other (п.5.5 matches п.5.5.1 and vice versa)
    """
    a = normalize_ref(ref_a)
    b = normalize_ref(ref_b)
    return a == b or a.startswith(b + ".") or b.startswith(a + ".")


def compute_metrics(
    gt_findings: list[dict],
    solomon_findings: list[dict],
    contract_clauses: list[dict],
) -> dict:
    """
    §13.2 metrics (grounding excluded — corpus not yet seeded).

    Precision  = TP / len(solomon_findings)
    Recall     = TP / len(gt_findings)
    Clause-ref accuracy = clauses_found_in_contract / len(solomon_findings)
    """
    gt_refs = [f["clause_ref"] for f in gt_findings]
    contract_refs = {normalize_ref(c["ref"]) for c in contract_clauses}

    tp_precision = 0  # Solomon findings that match a GT finding
    tp_recall_set: set = set()  # GT refs that Solomon found
    clause_ref_valid = 0  # Solomon findings whose clause_ref is in the contract

    per_finding = []
    for sf in solomon_findings:
        s_ref = normalize_ref(sf["clause_ref"])

        # Clause-ref accuracy: is the ref locatable in contract?
        in_contract = any(_refs_match(sf["clause_ref"], c["ref"]) for c in contract_clauses)
        if in_contract:
            clause_ref_valid += 1

        # Precision: does GT contain this ref?
        gt_match = next(
            (g for g in gt_findings if _refs_match(sf["clause_ref"], g["clause_ref"])),
            None,
        )
        is_tp = gt_match is not None
        if is_tp:
            tp_precision += 1
            tp_recall_set.add(normalize_ref(gt_match["clause_ref"]))

        per_finding.append({
            "clause_ref": sf["clause_ref"],
            "category": sf.get("category", "other"),
            "severity": sf.get("severity", "?"),
            "in_contract": in_contract,
            "gt_match": gt_match["clause_ref"] if gt_match else None,
            "is_tp": is_tp,
        })

    n_solomon = len(solomon_findings)
    n_gt = len(gt_findings)

    precision = tp_precision / n_solomon if n_solomon > 0 else 0.0
    recall = len(tp_recall_set) / n_gt if n_gt > 0 else 0.0
    clause_acc = clause_ref_valid / n_solomon if n_solomon > 0 else 0.0

    missed_gt = [
        g["clause_ref"] for g in gt_findings
        if normalize_ref(g["clause_ref"]) not in tp_recall_set
    ]

    return {
        "targets": {
            "precision": {"value": round(precision, 4), "target": PRECISION_TARGET, "pass": precision >= PRECISION_TARGET},
            "recall":    {"value": round(recall, 4),    "target": RECALL_TARGET,    "pass": recall >= RECALL_TARGET},
            "clause_ref_accuracy": {"value": round(clause_acc, 4), "target": CLAUSE_ACC_TARGET, "pass": clause_acc >= CLAUSE_ACC_TARGET},
        },
        "counts": {
            "solomon_findings": n_solomon,
            "gt_findings": n_gt,
            "true_positives": tp_precision,
            "false_positives": n_solomon - tp_precision,
            "missed_by_solomon": n_gt - len(tp_recall_set),
            "rejected_by_guardrail": 0,  # filled in by caller
            "clause_refs_invalid": n_solomon - clause_ref_valid,
        },
        "missed_gt_refs": missed_gt,
        "per_finding": per_finding,
    }


# ─── Output ──────────────────────────────────────────────────────────────────

PASS = "✅ PASS"
FAIL = "❌ FAIL"
WARN = "⚠️  WARN"


def print_results(metrics: dict, gt_findings: list[dict], solomon_findings: list[dict]):
    t = metrics["targets"]
    c = metrics["counts"]

    print()
    print("═" * 60)
    print("  SOLOMON CONTRACTS — PHASE GATE EVAL (§13.2)")
    print("  Gold set: Levays supply contract bundle")
    print("  Note: Grounding rate skipped — corpus not yet seeded")
    print("═" * 60)
    print(f"  Ground truth findings  : {c['gt_findings']}")
    print(f"  Solomon findings       : {c['solomon_findings']}")
    print(f"  Rejected by guardrail  : {c['rejected_by_guardrail']}")
    print()
    print("  ── Metrics ──────────────────────────────────")

    for key, label in [("precision", "Precision"), ("recall", "Recall"), ("clause_ref_accuracy", "Clause-ref accuracy")]:
        m = t[key]
        badge = PASS if m["pass"] else FAIL
        print(f"  {label:<22} {m['value']:.3f}  (target ≥ {m['target']})  {badge}")

    all_pass = all(m["pass"] for m in t.values())
    print()
    print(f"  ── Phase gate: {'PASSED ✅' if all_pass else 'NOT PASSED ❌ — iterate prompts before Phase 2 rollout'}")

    if metrics["missed_gt_refs"]:
        print()
        print("  ── Missed by Solomon (ground-truth refs not found):")
        for ref in metrics["missed_gt_refs"][:10]:
            print(f"     - {ref}")
        if len(metrics["missed_gt_refs"]) > 10:
            print(f"     ... and {len(metrics['missed_gt_refs']) - 10} more")

    fps = [f for f in metrics["per_finding"] if not f["is_tp"]]
    if fps:
        print()
        print("  ── False positives (Solomon refs not in ground truth):")
        for f in fps[:10]:
            print(f"     - {f['clause_ref']} [{f['category']}] sev={f['severity']}")
        if len(fps) > 10:
            print(f"     ... and {len(fps) - 10} more")

    print("═" * 60)
    print()


def save_results(
    metrics: dict,
    gt_findings: list[dict],
    solomon_findings: list[dict],
    rejected_count: int,
):
    metrics["counts"]["rejected_by_guardrail"] = rejected_count
    metrics["metadata"] = {
        "eval_timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "gold_set": "Levays",
        "grounding_measured": False,
        "note": "Grounding rate skipped — legal corpus not yet seeded. Re-run after task #7.",
    }
    RESULTS_FILE.write_text(json.dumps(metrics, ensure_ascii=False, indent=2))
    logger.info(f"Results saved to {RESULTS_FILE}")


# ─── Entry point ─────────────────────────────────────────────────────────────

def main():
    print()
    logger.info("Solomon Contracts — automated eval (§13)")
    logger.info(f"Looking for Levays files in: {LEVAYS_DIR}")

    # Validate input files
    contract_path = LEVAYS_DIR / EXPECTED_CONTRACT
    risk_note_path = LEVAYS_DIR / EXPECTED_RISK_NOTE

    missing = []
    if not contract_path.exists():
        missing.append(EXPECTED_CONTRACT)
    if not risk_note_path.exists():
        missing.append(EXPECTED_RISK_NOTE)

    if missing:
        print()
        print("ERROR: Levays gold-set files not found in training_data/levays/")
        print()
        print("Required files:")
        for name in [EXPECTED_CONTRACT, EXPECTED_RISK_NOTE]:
            status = "✅" if (LEVAYS_DIR / name).exists() else "❌ MISSING"
            print(f"  {status}  {name}")
        print()
        print("Place the Levays bundle in training_data/levays/ and re-run.")
        print("See training_data/levays/README.md for details.")
        sys.exit(2)

    # Step 1: Parse ground truth
    logger.info(f"Parsing ground truth from: {risk_note_path.name}")
    gt_findings = parse_risk_note(risk_note_path)
    if not gt_findings:
        logger.error("Ground truth parser returned 0 findings — check the risk note format")
        sys.exit(1)

    # Step 2: Run Solomon analyzer
    solomon_findings, contract_clauses = run_solomon_on_contract(contract_path)
    if not solomon_findings and not contract_clauses:
        logger.error("Analyzer returned no findings and no clauses — check contract extraction")
        sys.exit(1)

    # Step 3: Compute metrics
    rejected_count = 0  # guardrail rejects are counted inside scan_document
    metrics = compute_metrics(gt_findings, solomon_findings, contract_clauses)

    # Step 4: Output
    print_results(metrics, gt_findings, solomon_findings)
    save_results(metrics, gt_findings, solomon_findings, rejected_count)

    # Exit code: 0 = all pass, 1 = some failed
    all_pass = all(m["pass"] for m in metrics["targets"].values())
    sys.exit(0 if all_pass else 1)


if __name__ == "__main__":
    main()
